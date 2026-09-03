import io
import os

import docx
from pypdf import PdfReader

ALLOWED_EXTS = {".txt", ".md", ".pdf", ".docx"}

ATTACH_MAX_FILES = int(os.environ.get("ATTACH_MAX_FILES", "5"))
ATTACH_MAX_BYTES = int(os.environ.get("ATTACH_MAX_BYTES", "20971520"))
ATTACH_MAX_CHARS = int(os.environ.get("ATTACH_MAX_CHARS", "12000"))

DOC_MESSAGE = "legacy .doc not supported — convert to .docx"
UNSUPPORTED_MESSAGE = "unsupported file type: {ext}. Allowed: .md, .pdf, .txt, .docx"
EMPTY_MESSAGE = "file appears empty (image-only file? text extraction only)"
SCANNED_MESSAGE = "scanned PDF? OCR not supported yet"

MIN_CHARS = 50


def clean_text(s: str) -> str:
    return s.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")


def _sanitize_name(name: str) -> str:
    name = (name or "").replace("\x00", "").replace("\\", "/")
    return name.split("/")[-1].strip()[:200]


def truncate_head_tail(text: str, cap: int) -> str:
    if len(text) <= cap:
        return text
    head = cap * 2 // 3
    tail = max(0, cap - head - 100)
    marker = (
        f"\n\n[TRUNCATED — showing first {head} and last {tail}"
        f" of {len(text)} chars]\n\n"
    )
    out = text[:head] + marker
    if tail:
        out += text[-tail:]
    return out if len(out) <= cap else out[:cap]


def _extract_pdf(data: bytes) -> tuple[str, list[int]]:
    try:
        reader = PdfReader(io.BytesIO(data))
        raw_pages = [page.extract_text() or "" for page in reader.pages]
    except ValueError:
        raise
    except Exception:
        raise ValueError("not a valid .pdf file")
    pages_without_text = [
        i for i, t in enumerate(raw_pages, start=1) if len(t.strip()) < 1
    ]
    if sum(len(t.strip()) for t in raw_pages) < MIN_CHARS:
        raise ValueError(SCANNED_MESSAGE)
    return "\n\n".join(t for t in raw_pages if t.strip()), pages_without_text


def _extract_docx(data: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(data))
    except ValueError:
        raise
    except Exception:
        raise ValueError("not a valid .docx file")
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n\n".join(parts)


def _ext_error(ext: str) -> ValueError:
    return ValueError(UNSUPPORTED_MESSAGE.format(ext=ext or "(none)"))


def _split_ext(name: str) -> str:
    return os.path.splitext(name)[1].lower()


def extract_upload(filename: str, data: bytes) -> dict:
    name = _sanitize_name(filename)
    ext = _split_ext(name)
    if ext == ".doc":
        raise ValueError(DOC_MESSAGE)
    if ext not in ALLOWED_EXTS:
        raise _ext_error(ext)

    pages_without_text: list[int] | None = None
    if ext == ".pdf":
        text, pages_without_text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    else:
        text = data.decode("utf-8", errors="replace")

    text = clean_text(text)
    if len(text.strip()) == 0:
        raise ValueError(EMPTY_MESSAGE)
    text = truncate_head_tail(text, ATTACH_MAX_CHARS)

    result = {
        "name": name,
        "ext": ext,
        "size": len(data),
        "chars": len(text),
        "text": text,
    }
    if pages_without_text is not None:
        result["pages_without_text"] = pages_without_text
    return result


def validate_docs(docs: list[dict] | None) -> list[dict] | None:
    if not docs:
        return None
    if len(docs) > ATTACH_MAX_FILES:
        raise ValueError(f"too many attachments (max {ATTACH_MAX_FILES})")
    out: list[dict] = []
    for d in docs:
        name = _sanitize_name(d.get("name") or "")
        ext = _split_ext(name)
        if ext == ".doc":
            raise ValueError(DOC_MESSAGE)
        if ext not in ALLOWED_EXTS:
            raise _ext_error(ext)
        text = truncate_head_tail(clean_text(d.get("text") or ""), ATTACH_MAX_CHARS)
        out.append(
            {
                "name": name,
                "ext": ext,
                "size": int(d.get("size") or 0),
                "chars": len(text),
                "text": text,
            }
        )
    return out
